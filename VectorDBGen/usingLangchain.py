"""
Build a local (open-source, no login) Chroma vector DB from .docx and .toml files
using Azure OpenAI embeddings via LangChain.

Prereqs:
  pip install "langchain>=0.2.11" "langchain-openai>=0.1.18" \
              "langchain-chroma>=0.1.2" "langchain-text-splitters>=0.2.2" \
              chromadb python-docx tomli

Env vars (example):
  export KB_ROOT="/path/to/knowledge_base"
  export CHROMA_DIR="./chroma_kb"
  export COLLECTION_NAME="kb_chunks"

  export AZURE_OPENAI_API_KEY="<your-key>"
  export AZURE_OPENAI_ENDPOINT="https://<your-resource>.openai.azure.com"
  export AZURE_OPENAI_API_VERSION="2024-10-21"
  export AZURE_OPENAI_EMBEDDINGS_DEPLOYMENT="text-embeddings-002"  # your deployment name

Optional:
  export CHUNK_SIZE=1200
  export CHUNK_OVERLAP=200
  export BATCH_SIZE=256
"""

import os
import json
import hashlib
from pathlib import Path
from typing import List, Dict
from collections import defaultdict

# --- TOML parsing (stdlib in Py 3.11+, fallback to tomli)
try:
    import tomllib  # Python 3.11+
except ModuleNotFoundError:
    import tomli as tomllib

from docx import Document

# --- LangChain bits
from langchain_openai import AzureOpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LCDocument
from langchain_chroma import Chroma


# =========================
# Config
# =========================
KB_ROOT = Path(os.getenv("KB_ROOT", ".")).expanduser().resolve()
CHROMA_DIR = os.getenv("CHROMA_DIR", "./chroma_kb")
COLLECTION_NAME = os.getenv("COLLECTION_NAME", "kb_chunks")

# Azure OpenAI (LangChain's AzureOpenAIEmbeddings reads these explicitly)
AZURE_OPENAI_API_KEY = os.environ["AZURE_OPENAI_API_KEY"]
AZURE_OPENAI_ENDPOINT = os.environ["AZURE_OPENAI_ENDPOINT"]
AZURE_OPENAI_API_VERSION = os.getenv("AZURE_OPENAI_API_VERSION", "2024-10-21")
AZURE_OPENAI_EMBEDDINGS_DEPLOYMENT = os.environ["AZURE_OPENAI_EMBEDDINGS_DEPLOYMENT"]

CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "1200"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "200"))
BATCH_SIZE = int(os.getenv("BATCH_SIZE", "256"))


# =========================
# File loading
# =========================
def read_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts: List[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    # include simple table text (row-wise)
    for table in doc.tables:
        for row in table.rows:
            cells = [((c.text or "").strip()) for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def read_toml_text(path: Path) -> str:
    with open(path, "rb") as f:
        data = tomllib.load(f)
    # flatten to a stable JSON string for embeddings
    return json.dumps(data, ensure_ascii=False, indent=2)


def load_files_as_documents(root: Path) -> List[LCDocument]:
    """
    Return base (un-chunked) LangChain Documents with metadata {path, filename}.
    """
    docs: List[LCDocument] = []
    for p in root.rglob("*"):
        if not p.is_file():
            continue
        try:
            if p.suffix.lower() == ".docx":
                text = read_docx_text(p)
            elif p.suffix.lower() == ".toml":
                text = read_toml_text(p)
            else:
                continue
            if not text:
                continue
            docs.append(
                LCDocument(
                    page_content=text,
                    metadata={
                        "path": str(p),
                        "filename": p.name,
                    },
                )
            )
        except Exception as e:
            print(f"[WARN] Failed to read {p}: {e}")
    return docs


# =========================
# Helpers
# =========================
def deterministic_id(*parts) -> str:
    h = hashlib.sha256()
    for part in parts:
        h.update(str(part).encode("utf-8"))
    return h.hexdigest()  # safe as Chroma id


def add_chunk_ids(chunk_docs: List[LCDocument]) -> None:
    """
    Add a sequential chunk_id per original file path into metadata.
    """
    counters: Dict[str, int] = defaultdict(int)
    for d in chunk_docs:
        path = d.metadata.get("path", "unknown")
        d.metadata["chunk_id"] = str(counters[path])
        counters[path] += 1


# =========================
# Main build
# =========================
def main():
    # 1) Load files into base Documents
    base_docs = load_files_as_documents(KB_ROOT)
    if not base_docs:
        print(f"[INFO] No .docx or .toml files found under {KB_ROOT}")
        return

    # 2) Chunk with LangChain splitter (propagates metadata)
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
    )
    chunk_docs = splitter.split_documents(base_docs)
    add_chunk_ids(chunk_docs)

    # 3) Configure Azure OpenAI embeddings (LangChain)
    embeddings = AzureOpenAIEmbeddings(
        api_key=AZURE_OPENAI_API_KEY,
        azure_endpoint=AZURE_OPENAI_ENDPOINT,
        azure_deployment=AZURE_OPENAI_EMBEDDINGS_DEPLOYMENT,  # deployment name
        openai_api_version=AZURE_OPENAI_API_VERSION,
    )

    # 4) Create/get a persistent Chroma collection and add chunks in batches
    vectorstore = Chroma(
        collection_name=COLLECTION_NAME,
        persist_directory=CHROMA_DIR,
        embedding_function=embeddings,
    )

    # Batch add for throughput
    for i in range(0, len(chunk_docs), BATCH_SIZE):
        batch = chunk_docs[i : i + BATCH_SIZE]
        ids = [
            deterministic_id(d.metadata.get("path", ""), d.metadata.get("chunk_id", ""))
            for d in batch
        ]
        vectorstore.add_documents(documents=batch, ids=ids)
        print(f"[OK] Upserted {len(batch)} chunks (running total: {min(i + BATCH_SIZE, len(chunk_docs))})")

    # Ensure data flushed to disk
    vectorstore.persist()
    print(f"[DONE] Vector DB ready in Chroma (dir='{CHROMA_DIR}', collection='{COLLECTION_NAME}'). Total chunks: {len(chunk_docs)}")


if __name__ == "__main__":
    main()
