"""
Create a local, persistent vector DB (Chroma) from a knowledge base of .docx and .toml files
using Azure OpenAI text-embeddings-002. No login/userId/password needed for Chroma.

Prereqs:
  pip install "openai>=1.40.0" chromadb python-docx tomli

Env vars (example):
  export KB_ROOT="/path/to/your/knowledge_base"
  export CHROMA_DIR="./chroma_kb"

  export AZURE_OPENAI_ENDPOINT="https://<your-resource>.openai.azure.com"
  export AZURE_OPENAI_API_KEY="<your-key>"
  export AZURE_OPENAI_API_VERSION="2024-10-21"
  export AZURE_OPENAI_EMBEDDINGS_DEPLOYMENT="text-embeddings-002"
"""

import os
import json
import time
import hashlib
from pathlib import Path
from typing import List, Tuple

# TOML parsing (Python 3.11+ has tomllib; fallback to tomli for older Pythons)
try:
    import tomllib  # py>=3.11
except ModuleNotFoundError:
    import tomli as tomllib  # fallback for py<3.11

from docx import Document

# Azure OpenAI SDK
from openai import AzureOpenAI

# ChromaDB (local, persistent)
import chromadb
from chromadb.config import Settings

# =========================
# Config
# =========================
KB_ROOT = Path(os.getenv("KB_ROOT", ".")).expanduser().resolve()
CHROMA_DIR = os.getenv("CHROMA_DIR", "./chroma_kb")
COLLECTION_NAME = os.getenv("COLLECTION_NAME", "kb_chunks")

AZURE_OPENAI_ENDPOINT = os.environ["AZURE_OPENAI_ENDPOINT"]
AZURE_OPENAI_API_KEY = os.environ["AZURE_OPENAI_API_KEY"]
AZURE_OPENAI_API_VERSION = os.getenv("AZURE_OPENAI_API_VERSION", "2024-10-21")
EMBED_DEPLOYMENT = os.environ["AZURE_OPENAI_EMBEDDINGS_DEPLOYMENT"]  # e.g., "text-embeddings-002"

CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "1200"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "200"))

# =========================
# File loading
# =========================
def read_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    # tables (keep it simple, row-wise)
    for table in doc.tables:
        for row in table.rows:
            cells = [((c.text or "").strip()) for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()

def read_toml_text(path: Path) -> str:
    with open(path, "rb") as f:
        data = tomllib.load(f)
    # Stable, human-readable representation
    return json.dumps(data, ensure_ascii=False, indent=2)

def walk_kb(root: Path) -> List[Tuple[str, str]]:
    """
    Return list of (absolute_path, text_content) for .docx and .toml files.
    """
    items: List[Tuple[str, str]] = []
    for p in root.rglob("*"):
        if not p.is_file():
            continue
        try:
            suf = p.suffix.lower()
            if suf == ".docx":
                items.append((str(p), read_docx_text(p)))
            elif suf == ".toml":
                items.append((str(p), read_toml_text(p)))
        except Exception as e:
            print(f"[WARN] Failed to read {p}: {e}")
    return items

# =========================
# Chunking
# =========================
def chunk_text(text: str, chunk_chars: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    if not text:
        return []
    n = len(text)
    out: List[str] = []
    start = 0
    while start < n:
        end = min(n, start + chunk_chars)
        out.append(text[start:end])
        if end == n:
            break
        start = max(end - overlap, start + 1)
    return out

# =========================
# Deterministic IDs
# =========================
def deterministic_id(*parts) -> str:
    h = hashlib.sha256()
    for part in parts:
        h.update(str(part).encode("utf-8"))
    return h.hexdigest()  # valid Chroma id

# =========================
# Azure OpenAI embeddings (batched + retries)
# =========================
aoai = AzureOpenAI(
    api_key=AZURE_OPENAI_API_KEY,
    api_version=AZURE_OPENAI_API_VERSION,
    azure_endpoint=AZURE_OPENAI_ENDPOINT,
)

def embed_texts(texts: List[str], batch: int = 32, max_retries: int = 5) -> List[List[float]]:
    vectors: List[List[float]] = []
    for i in range(0, len(texts), batch):
        sub = texts[i:i + batch]
        for attempt in range(1, max_retries + 1):
            try:
                resp = aoai.embeddings.create(model=EMBED_DEPLOYMENT, input=sub)
                vectors.extend([d.embedding for d in resp.data])
                break
            except Exception as e:
                if attempt == max_retries:
                    raise
                wait_s = min(2 ** attempt, 20)
                print(f"[WARN] Embeddings batch failed (attempt {attempt}): {e} → retrying in {wait_s}s")
                time.sleep(wait_s)
    return vectors

# =========================
# Chroma helpers
# =========================
def get_collection() -> "chromadb.api.models.Collection.Collection":
    # Persistent client: stores data on disk in CHROMA_DIR
    client = chromadb.PersistentClient(
        path=CHROMA_DIR,
        settings=Settings(anonymized_telemetry=False),
    )
    # Create or load collection; cosine distance is default; set hnsw space if desired
    coll = client.get_or_create_collection(
        name=COLLECTION_NAME,
        metadata={"hnsw:space": "cosine"}
    )
    return coll

def add_batch(
    coll,
    ids: List[str],
    documents: List[str],
    metadatas: List[dict],
    embeddings: List[List[float]],
):
    # Chroma expects parallel lists with equal lengths
    coll.add(ids=ids, documents=documents, metadatas=metadatas, embeddings=embeddings)

# =========================
# Main
# =========================
def main():
    # 1) Load files & chunk
    file_texts = walk_kb(KB_ROOT)
    records: List[Tuple[str, str, str, int]] = []  # (abs_path, filename, chunk_text, chunk_idx)
    for abs_path, text in file_texts:
        if not text:
            continue
        chunks = chunk_text(text)
        fname = Path(abs_path).name
        for i, ch in enumerate(chunks):
            records.append((abs_path, fname, ch, i))

    if not records:
        print(f"[INFO] No .docx or .toml files found under {KB_ROOT}")
        return

    # 2) Create/load Chroma collection (local persistent)
    coll = get_collection()

    # 3) Embed & add in batches
    BATCH = 256
    total = 0
    for i in range(0, len(records), BATCH):
        batch = records[i:i + BATCH]

        docs = [r[2] for r in batch]
        vecs = embed_texts(docs)

        ids = [deterministic_id(r[0], r[3]) for r in batch]
        metadatas = [
            {
                "path": r[0],
                "filename": r[1],
                "chunk_id": str(r[3]),
            }
            for r in batch
        ]

        add_batch(
            coll=coll,
            ids=ids,
            documents=docs,
            metadatas=metadatas,
            embeddings=vecs,
        )
        total += len(batch)
        print(f"[OK] Upserted {len(batch)} chunks (running total: {total})")

    print(f"[DONE] Vector DB created & populated in Chroma (dir='{CHROMA_DIR}', collection='{COLLECTION_NAME}'). Total chunks: {total}")

if __name__ == "__main__":
    main()
